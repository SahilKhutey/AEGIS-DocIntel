"""
AEGIS-AMDI-OS — DOCX Loader
=============================
Microsoft Word document loader using python-docx.
"""
from __future__ import annotations

import io
import logging
from typing import Any

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from src.core.document_object import DocumentFormat, DocumentObject
from src.ingestion.base import BaseLoader, FormatError

logger = logging.getLogger(__name__)


class DOCXLoader(BaseLoader):
    """DOCX document loader."""

    FORMAT_NAME = "docx"
    SUPPORTED_EXTENSIONS = {".docx"}
    DOCX_MAGIC = b"PK\x03\x04"  # ZIP-based format

    def __init__(self, max_size_mb: int = 100, **options):
        super().__init__(**options)
        self.max_size_mb = max_size_mb

    def validate(self, raw_bytes: bytes) -> bool:
        """Check if bytes are a valid DOCX."""
        if not raw_bytes or len(raw_bytes) < 4:
            return False
        # DOCX is a ZIP file with specific structure
        if raw_bytes[:4] != self.DOCX_MAGIC:
            return False
        return True

    async def load(self, source, filename: str = "") -> DocumentObject:
        """Load a DOCX document."""
        raw_bytes, name = self.read_source(source)
        if filename:
            name = filename
        if not name:
            name = "document.docx"

        if not self.validate(raw_bytes):
            raise FormatError("Not a valid DOCX file")

        size_mb = len(raw_bytes) / (1024 * 1024)
        if size_mb > self.max_size_mb:
            from src.ingestion.base import SizeLimitError
            raise SizeLimitError(f"DOCX too large: {size_mb:.1f}MB")

        metadata = self._extract_metadata(raw_bytes)
        # Build text content
        text_parts = []
        try:
            doc = DocxDocument(io.BytesIO(raw_bytes))
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            # Also extract table text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
        except Exception as e:
            logger.warning(f"Text extraction failed: {e}")

        doc_obj = DocumentObject(
            filename=name,
            format=DocumentFormat.DOCX,
            raw_bytes=raw_bytes,
            metadata=metadata,
            word_count=sum(len(t.split()) for t in text_parts),
            text_content="\n\n".join(text_parts),
        )
        return doc_obj

    def _extract_metadata(self, raw_bytes: bytes) -> dict[str, Any]:
        """Extract DOCX metadata."""
        metadata: dict[str, Any] = {}
        try:
            doc = DocxDocument(io.BytesIO(raw_bytes))
            cp = doc.core_properties
            metadata["title"] = cp.title
            metadata["author"] = cp.author
            metadata["subject"] = cp.subject
            metadata["keywords"] = cp.keywords
            metadata["created"] = str(cp.created) if cp.created else None
            metadata["modified"] = str(cp.modified) if cp.modified else None
            metadata["last_modified_by"] = cp.last_modified_by
            metadata["revision"] = cp.revision

            # Count elements
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)
            metadata["section_count"] = len(doc.sections)
            # Count images
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    image_count += 1
            metadata["image_count"] = image_count
        except Exception as e:
            logger.warning(f"DOCX metadata extraction failed: {e}")
        return metadata
